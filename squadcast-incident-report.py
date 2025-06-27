import pandas as pd
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
from datetime import datetime
import re

# --- Date formatting for filenames ---
month_year = datetime.now().strftime("%B-%Y").lower()  # e.g., june-2025
chart_filename = f"alerts-by-service-{month_year}.png"
doc_filename = f"squadcast-incident-review-report-{month_year}.docx"

# --- Load and clean CSV ---
df = pd.read_csv("incident_report.csv")  # Replace with your actual file

# --- Clean ID if formatted as HYPERLINK in Excel exports ---
def extract_id(cell):
    match = re.search(r'https:\/\/app\.squadcast\.com\/incident\/([a-zA-Z0-9]+)', str(cell))
    return match.group(1) if match else str(cell)

df['id'] = df['id'].apply(extract_id)

# --- Preprocessing ---
df['ttr_minutes'] = df['ttr (ms)'].fillna(0) / 60000
df['title'] = df['title'].fillna("Untitled")
df['service'] = df['service'].fillna("Unknown")

# --- Summaries ---
total_alerts = len(df)
alerts_per_service = df['service'].value_counts()
top_alerts = df['title'].value_counts()
top_ttr_alerts = df.sort_values(by='ttr_minutes', ascending=False).head(5)
long_ttr_alerts = df[df['ttr_minutes'] > 60]

# --- Save Chart ---
plt.figure(figsize=(6, 4))
alerts_per_service.plot(kind='bar', color='skyblue')
plt.title("Alerts per Service")
plt.xlabel("Service")
plt.ylabel("Alert Count")
plt.xticks(rotation=45, ha='right')
plt.tight_layout()
plt.savefig(chart_filename)
plt.close()

# --- Create Word Report ---
doc = Document()
doc.add_heading('Incident Review Report', 0)

doc.add_heading('1. Total Alert Count', level=1)
doc.add_paragraph(f"Total number of alerts: {total_alerts}")

doc.add_heading('2. Alerts per Service', level=1)
table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid'
table.rows[0].cells[0].text = 'Service'
table.rows[0].cells[1].text = 'Alert Count'
for service, count in alerts_per_service.items():
    row = table.add_row().cells
    row[0].text = service
    row[1].text = str(count)

doc.add_picture(chart_filename, width=Inches(5.5))

doc.add_heading('3. Most Frequently Triggered Alerts', level=1)
table = doc.add_table(rows=1, cols=3)
table.style = 'Light Grid'
table.rows[0].cells[0].text = 'Alert Title'
table.rows[0].cells[1].text = 'Count'
table.rows[0].cells[2].text = 'Sample Alert URL'
for title, count in top_alerts.items():
    last_id = df[df['title'] == title]['id'].iloc[-1]
    incident_url = f"https://app.squadcast.com/incident/{last_id}"
    row = table.add_row().cells
    row[0].text = title
    row[1].text = str(count)
    row[2].text = incident_url

doc.add_heading('4. Alerts with the Highest Time to Resolve', level=1)
table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid'
table.rows[0].cells[0].text = 'Title'
table.rows[0].cells[1].text = 'Service'
table.rows[0].cells[2].text = 'Time to Resolve'
table.rows[0].cells[3].text = 'Alert ID'
for _, row in top_ttr_alerts.iterrows():
    ttr = row['ttr_minutes']
    if ttr >= 60:
        hours = int(ttr // 60)
        minutes = int(ttr % 60)
        ttr_formatted = f"{hours}h {minutes}m"
    else:
        ttr_formatted = f"{int(ttr)} mins"
    cells = table.add_row().cells
    cells[0].text = row['title']
    cells[1].text = row['service']
    cells[2].text = ttr_formatted
    cells[3].text = row['id']

doc.add_heading('5. Alerts Taking More Than 1 Hour to Resolve', level=1)
table = doc.add_table(rows=1, cols=4)
table.style = 'Light Grid'
table.rows[0].cells[0].text = 'Title'
table.rows[0].cells[1].text = 'Service'
table.rows[0].cells[2].text = 'Time to Resolve'
table.rows[0].cells[3].text = 'Alert ID'
for _, row in long_ttr_alerts.iterrows():
    ttr = row['ttr_minutes']
    hours = int(ttr // 60)
    minutes = int(ttr % 60)
    ttr_formatted = f"{hours}h {minutes}m"
    cells = table.add_row().cells
    cells[0].text = row['title']
    cells[1].text = row['service']
    cells[2].text = ttr_formatted
    cells[3].text = row['id']

# --- Save Word File ---
doc.save(doc_filename)
print(f"Report generated: {doc_filename}")
